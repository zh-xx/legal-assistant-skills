from __future__ import annotations

import argparse
import shutil
import subprocess
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse

import pdfplumber
from docx import Document

from cn_law_normalizer import normalize_cn_law_markdown
from stage3_checker import run_stage3_checks, write_stage3_report_md

MINERU_OCR_INSTALL_URL = "https://github.com/cat-xierluo/legal-skills/tree/main/skills/mineru-ocr"
ARTIFACT_LEVELS = ("minimal", "standard", "debug")


def _write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def _derive_stage_paths(base_out_path: Path) -> tuple[Path, Path]:
    suffix = base_out_path.suffix or ".md"
    stem = base_out_path.stem
    stage1 = base_out_path.with_name(f"{stem}.stage1{suffix}")
    stage2 = base_out_path.with_name(f"{stem}.stage2{suffix}")
    return stage1, stage2


def _derive_stage3_report_path(stage2_path: Path) -> Path:
    name = stage2_path.name
    if name.endswith(".stage2.md"):
        return stage2_path.with_name(name.replace(".stage2.md", ".stage3-check.md"))
    return stage2_path.with_name(f"{stage2_path.stem}.stage3-check.md")


def _derive_legacy_stage3_txt_path(stage2_path: Path) -> Path:
    name = stage2_path.name
    if name.endswith(".stage2.md"):
        return stage2_path.with_name(name.replace(".stage2.md", ".stage3-check.txt"))
    return stage2_path.with_name(f"{stage2_path.stem}.stage3-check.txt")


def _derive_user_output_paths(base_dir: Path, input_stem: str) -> dict[str, Path]:
    return {
        "deliverable": base_dir / f"{input_stem}+最终成果.md",
        "report": base_dir / f"{input_stem}+审核报告.md",
    }


def _safe_unlink(path: Path | None) -> None:
    if path is None or not path.exists():
        return
    try:
        path.unlink()
    except OSError:
        pass


def _apply_artifact_policy(
    *,
    stage1_path: Path,
    stage2_path: Path,
    stage3_md_path: Path | None,
    deliverable_path: Path | None,
    report_path: Path,
    artifact_level: str,
    review_status: str,
) -> None:
    if artifact_level == "debug":
        return

    rejected = review_status in {"rejected_non_law", "rejected_check_failed"}

    if artifact_level == "minimal":
        keep: set[Path] = {report_path}
        if deliverable_path is not None and deliverable_path.exists():
            keep.add(deliverable_path)
        if rejected:
            keep.add(stage1_path)
            keep.add(stage2_path)
            if stage3_md_path is not None and stage3_md_path.exists():
                keep.add(stage3_md_path)
        candidates: list[Path | None] = [
            stage1_path,
            stage2_path,
            stage3_md_path,
            deliverable_path,
            report_path,
        ]
        for candidate in candidates:
            if candidate is None:
                continue
            if candidate in keep:
                continue
            _safe_unlink(candidate)
    return


def _write_review_report_md(
    report_path: Path,
    *,
    input_ref: str,
    stage1_path: Path,
    stage2_path: Path,
    stage3_md_path: Path | None,
    stage3_skipped: bool,
    stage3_pass: bool | None,
    stage3_attempts: int,
    stage3_attempt_details: list[dict[str, object]] | None,
    stage2_last_reason: str,
    law_decision: str,
    engine: str,
    review_status: str,
    deliverable_path: Path | None,
) -> None:
    status_label = {"PASS": "通过", "FAIL": "未通过"}
    stage3_status = "已跳过" if stage3_skipped else ("通过" if stage3_pass else "未通过")
    reason_map = {
        "applied": "已完成格式优化",
        "already-normalized": "已是目标格式，无需改动",
        "non-law-document": "识别为非法律文档，已按规则拒绝二阶段处理",
        "legal-structure-not-detected": "未识别法律结构，二阶段未应用",
        "preserve-check-failed": "保真校验未通过，已回退",
        "autofix": "已执行自动修复",
    }
    review_label = {
        "approved": "通过",
        "rejected_non_law": "拒绝（非法律文档）",
        "rejected_check_failed": "拒绝（检查未通过）",
    }.get(review_status, review_status)

    lines: list[str] = []
    lines.append("# 文档转换审核报告")
    lines.append("")
    lines.append("## 1. 文档基本信息")
    lines.append("")
    lines.append(f"生成时间：{datetime.now().isoformat(timespec='seconds')}")
    lines.append(f"输入文件：{input_ref}")
    lines.append(f"转换引擎：{engine}")
    decision_label = {
        "law": "法律文本",
        "non-law": "非法律文本",
        "auto": "自动判断（由规则或模型识别）",
    }.get(law_decision, f"自动判断（当前值：{law_decision}）")
    lines.append(f"文档类型判断：{decision_label}")
    lines.append(f"最终审核结论：{review_label}")
    lines.append("")

    lines.append("## 2. Skill 执行说明")
    lines.append("")
    lines.append("- 第一阶段：调用转换引擎将原文转为 Markdown。")
    lines.append("- 第二阶段：按法律结构规则执行格式整理（仅格式处理）。")
    lines.append("- 第三阶段：执行内容准确性与效果检测，输出检查结论。")
    lines.append("")

    lines.append("## 3. 第一阶段结论")
    lines.append("")
    lines.append(f"- 输出文件：{stage1_path}")
    lines.append(f"- 结果：{'已生成' if stage1_path.exists() else '未生成'}")
    lines.append("")

    lines.append("## 4. 第二阶段结论")
    lines.append("")
    lines.append(f"- 输出文件：{stage2_path}")
    lines.append(f"- 处理结论：{reason_map.get(stage2_last_reason, stage2_last_reason or '未知')}")
    lines.append("")

    lines.append("## 5. 第三阶段结论")
    lines.append("")
    lines.append(f"- 检查状态：{stage3_status}")
    lines.append(f"- 尝试次数：{stage3_attempts}")
    if stage3_md_path:
        lines.append(f"- 检查明细文件：{stage3_md_path}")
    if stage3_attempt_details:
        lines.append("")
        lines.append("### 逐次检查结果")
        lines.append("")
        for entry in stage3_attempt_details:
            attempt_no = int(entry.get("attempt", 0))
            auto_fix_flag = "（自动修复复检）" if entry.get("auto_fix_applied") else ""
            lines.append(
                f"- 第 {attempt_no + 1} 次检查{auto_fix_flag}："
                f"A={'通过' if entry.get('stage3a_pass') else '未通过'}，"
                f"B={'通过' if entry.get('stage3b_pass') else '未通过'}，"
                f"综合={'通过' if entry.get('overall_pass') else '未通过'}"
            )
            for check in entry.get("checks", []):
                check_status = status_label.get(str(check.get("status", "")), str(check.get("status", "")))
                lines.append(f"  - {check.get('name', check.get('id', '检查项'))}：{check_status}；{check.get('detail', '')}")
    lines.append("")

    lines.append("## 6. 最终结论")
    lines.append("")
    lines.append(f"- 审核结论：{review_label}")
    lines.append(f"- 是否可交付：{'是' if review_status == 'approved' else '否'}")
    if review_status == "rejected_non_law":
        lines.append("- 拒绝原因：该文件判定为非法律文档，按策略拒绝产出最终成果。")
    elif review_status == "rejected_check_failed":
        lines.append("- 拒绝原因：第三阶段检查未通过。")
    lines.append("")

    lines.append("## 7. 交付物清单")
    lines.append("")
    if deliverable_path and deliverable_path.exists():
        lines.append(f"- 最终成果：{deliverable_path}")
    else:
        lines.append("- 最终成果：无（本次不予交付）")
    lines.append(f"- 审核报告：{report_path}")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _finalize_conversion(
    *,
    input_ref: str,
    input_stem: str,
    stage1_path: Path,
    stage2_path: Path,
    pipeline: dict[str, object],
    law_decision: str,
    engine: str,
    artifact_level: str,
    stage3_strict: bool,
    stage1_log_label: str,
    print_engine: str | None = None,
) -> None:
    review_status = str(pipeline.get("review_status", "rejected_check_failed"))
    user_outputs = _derive_user_output_paths(stage1_path.parent, input_stem)
    deliverable_path = user_outputs["deliverable"]
    report_path = user_outputs["report"]

    if review_status == "approved":
        shutil.copyfile(stage2_path, deliverable_path)
    else:
        _safe_unlink(deliverable_path)

    _write_review_report_md(
        report_path,
        input_ref=input_ref,
        stage1_path=stage1_path,
        stage2_path=stage2_path,
        stage3_md_path=pipeline.get("stage3_md_report"),
        stage3_skipped=bool(pipeline.get("stage3_skipped")),
        stage3_pass=pipeline.get("stage3_pass"),
        stage3_attempts=int(pipeline.get("stage3_attempts", 0)),
        stage3_attempt_details=pipeline.get("stage3_attempt_details"),
        stage2_last_reason=str(pipeline.get("stage2_last_reason", "")),
        law_decision=law_decision,
        engine=engine,
        review_status=review_status,
        deliverable_path=deliverable_path if deliverable_path.exists() else None,
    )
    _apply_artifact_policy(
        stage1_path=stage1_path,
        stage2_path=stage2_path,
        stage3_md_path=pipeline.get("stage3_md_report"),
        deliverable_path=deliverable_path if deliverable_path.exists() else None,
        report_path=report_path,
        artifact_level=artifact_level,
        review_status=review_status,
    )

    print(f"Saved review report: {report_path}")
    if deliverable_path.exists():
        print(f"Saved deliverable: {deliverable_path}")
    else:
        print("Saved deliverable: none (rejected)")
    print(f"Result: {review_status.upper()}")
    if review_status == "rejected_check_failed" and stage3_strict:
        raise SystemExit(f"Stage3 check failed. See review report: {report_path}")
    if stage1_path.exists():
        print(f"{stage1_log_label}: {stage1_path}")
    if stage2_path.exists():
        print(f"Saved stage2: {stage2_path}")
    if print_engine:
        print(print_engine)


def _prepare_stage2(stage1_path: Path, stage2_path: Path) -> None:
    stage2_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(stage1_path, stage2_path)


def _convert_txt(input_path: Path, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(input_path, out_path)


def _convert_docx(input_path: Path, out_path: Path) -> None:
    doc = Document(str(input_path))

    parts: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").rstrip()
        if text:
            parts.append(text)
        elif parts and parts[-1] != "":
            parts.append("")

    for table in doc.tables:
        parts.append("")
        for row in table.rows:
            row_text = "\t".join((cell.text or "").strip() for cell in row.cells).strip()
            if row_text:
                parts.append(row_text)
        parts.append("")

    text = "\n\n".join([p for p in parts if p is not None]).strip() + "\n"
    _write_text(out_path, text)


def _convert_pdf_pdfplumber(input_path: Path, out_path: Path) -> None:
    chunks: list[str] = []
    with pdfplumber.open(str(input_path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            chunks.append(f"<!-- Page {i} -->")
            page_text = page.extract_text() or ""
            chunks.append(page_text.rstrip())
            chunks.append("")
    text = "\n".join(chunks).rstrip() + "\n"
    _write_text(out_path, text)


def _resolve_mineru_ocr_convert_script() -> Path | None:
    candidates: list[Path] = [
        (Path.home() / ".codex" / "skills" / "mineru-ocr").resolve(),
        (Path.cwd() / ".claude" / "skills" / "mineru-ocr").resolve(),
    ]

    for root in candidates:
        script = root / "scripts" / "convert.js"
        if script.exists():
            return script
    return None


def _convert_with_mineru_ocr_skill(input_path: Path, out_path: Path) -> tuple[bool, str]:
    script = _resolve_mineru_ocr_convert_script()
    if script is None:
        return False, "mineru-ocr skill not found"

    cmd = ["/usr/bin/osascript", "-l", "JavaScript", str(script), str(input_path)]
    proc = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8", errors="replace")
    output = "\n".join([x for x in [proc.stdout.strip(), proc.stderr.strip()] if x])
    if proc.returncode != 0 or "转换失败" in output:
        return False, output or f"mineru-ocr exited with code {proc.returncode}"

    generated_md = input_path.with_suffix(".md")
    if not generated_md.exists():
        return False, f"mineru-ocr did not produce markdown: {generated_md}"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    if generated_md.resolve() != out_path.resolve():
        shutil.copyfile(generated_md, out_path)
        try:
            generated_md.unlink()
        except OSError:
            pass

    return True, output


def _run_stage2_normalize(out_path: Path, law_decision: str, stage2_profile: str = "default") -> dict[str, object]:
    text = out_path.read_text(encoding="utf-8")
    if law_decision in {"law", "non-law"}:
        print(f"Stage2 Classifier: {law_decision} (provided-by-caller)")
    else:
        print("Stage2 Classifier: fallback-rules (auto)")
    new_text, applied, stats = normalize_cn_law_markdown(
        text,
        law_decision=law_decision,
        stage2_profile=stage2_profile,
    )
    reason = stats.get("reason", "")
    if reason == "non-law-document":
        print("Stage2: rejected (non-law-document)")
        return {"applied": False, "reason": reason, "stats": stats}
    if not stats.get("preserve_check_passed", True):
        print("Stage2: skipped (preserve-check-failed)")
        return {"applied": False, "reason": "preserve-check-failed", "stats": stats}
    if not stats.get("legal_structure_detected", False):
        print("Stage2: no-op (legal-structure-not-detected)")
        return {"applied": False, "reason": "legal-structure-not-detected", "stats": stats}
    if applied:
        out_path.write_text(new_text, encoding="utf-8")
        print(
            "Stage2: applied "
            f"(title={stats.get('title_count', 0)}, "
            f"part={stats.get('part_count', 0)}, "
            f"chapter={stats.get('chapter_count', 0)}, "
            f"section={stats.get('section_count', 0)}, "
            f"article={stats.get('article_count', 0)}, "
            f"item_split={stats.get('item_split_count', 0)}, "
            f"space_cleanup={stats.get('space_cleanup_count', 0)})"
        )
        return {"applied": True, "reason": "applied", "stats": stats}
    print("Stage2: no-op (already-normalized)")
    return {"applied": False, "reason": "already-normalized", "stats": stats}


def _stage2_profile_for_attempt(attempt: int) -> str:
    if attempt <= 0:
        return "default"
    if attempt == 1:
        return "structure"
    return "minimal"


def _apply_stage3_autofix(stage2_path: Path, law_decision: str) -> bool:
    text = stage2_path.read_text(encoding="utf-8")
    new_text, applied, _ = normalize_cn_law_markdown(
        text,
        law_decision=law_decision,
        stage2_profile="default",
    )
    if applied:
        stage2_path.write_text(new_text, encoding="utf-8")
    return applied


def _run_stage2_stage3_pipeline(
    stage1_path: Path,
    stage2_path: Path,
    law_decision: str,
    skip_stage3_check: bool,
    stage3_max_retries: int,
) -> dict[str, object]:
    if skip_stage3_check:
        stage2_result = _run_stage2_normalize(stage2_path, law_decision=law_decision, stage2_profile="default")
        stage2_reason = str(stage2_result.get("reason", ""))
        review_status = "rejected_non_law" if stage2_reason == "non-law-document" else "approved"
        return {
            "stage3_skipped": True,
            "stage3_pass": None,
            "stage3_attempts": 0,
            "stage3_md_report": None,
            "stage3_attempt_details": [],
            "stage2_last_reason": stage2_reason,
            "review_status": review_status,
        }

    report_path = _derive_stage3_report_path(stage2_path)
    attempts: list[dict[str, object]] = []
    final_pass = False
    stage2_last_reason = ""
    review_status = "rejected_check_failed"

    for attempt in range(max(0, stage3_max_retries) + 1):
        profile = _stage2_profile_for_attempt(attempt)
        if attempt > 0:
            _prepare_stage2(stage1_path, stage2_path)
        print(f"Stage2 Retry: attempt={attempt} profile={profile}")
        stage2_result = _run_stage2_normalize(stage2_path, law_decision=law_decision, stage2_profile=profile)
        stage2_last_reason = str(stage2_result.get("reason", ""))
        check = run_stage3_checks(
            stage1_path=stage1_path,
            stage2_path=stage2_path,
            law_decision=law_decision,
            attempt=attempt,
            stage2_reason=stage2_last_reason,
        )
        check["stage2_profile"] = profile
        attempts.append(check)
        print(
            f"Stage3: attempt={attempt} "
            f"A={'PASS' if check.get('stage3a_pass') else 'FAIL'} "
            f"B={'PASS' if check.get('stage3b_pass') else 'FAIL'} "
            f"overall={'PASS' if check.get('overall_pass') else 'FAIL'}"
        )
        if check.get("business_decision") == "rejected_non_law":
            review_status = "rejected_non_law"
            break
        if check.get("overall_pass"):
            final_pass = True
            review_status = "approved"
            break

        if check.get("stage3b_auto_fixable_fail"):
            autofix_applied = _apply_stage3_autofix(stage2_path, law_decision=law_decision)
            if autofix_applied:
                autofix_check = run_stage3_checks(
                    stage1_path=stage1_path,
                    stage2_path=stage2_path,
                    law_decision=law_decision,
                    attempt=attempt,
                    stage2_reason="autofix",
                )
                stage2_last_reason = "autofix"
                autofix_check["stage2_profile"] = f"{profile}+autofix"
                autofix_check["auto_fix_applied"] = True
                attempts.append(autofix_check)
                print(
                    f"Stage3: attempt={attempt} autofix "
                    f"A={'PASS' if autofix_check.get('stage3a_pass') else 'FAIL'} "
                    f"B={'PASS' if autofix_check.get('stage3b_pass') else 'FAIL'} "
                    f"overall={'PASS' if autofix_check.get('overall_pass') else 'FAIL'}"
                )
                if autofix_check.get("business_decision") == "rejected_non_law":
                    review_status = "rejected_non_law"
                    break
                if autofix_check.get("overall_pass"):
                    final_pass = True
                    review_status = "approved"
                    break

    write_stage3_report_md(
        report_path,
        {
            "overall_pass": final_pass,
            "attempts": attempts,
        },
    )
    legacy_txt = _derive_legacy_stage3_txt_path(stage2_path)
    if legacy_txt.exists():
        legacy_txt.unlink()
    print(f"Saved stage3 report: {report_path}")
    return {
        "stage3_skipped": False,
        "stage3_pass": final_pass,
        "stage3_attempts": len(attempts),
        "stage3_md_report": report_path,
        "stage3_attempt_details": attempts,
        "stage2_last_reason": stage2_last_reason,
        "review_status": review_status,
    }


def _mineru_ocr_install_hint() -> str:
    return (
        "mineru-ocr skill 未安装。\n"
        f"- 安装地址: {MINERU_OCR_INSTALL_URL}\n"
        "- 安装命令:\n"
        "  python3 /Users/jicheng/.codex/skills/.system/skill-installer/scripts/install-skill-from-github.py "
        f"--url {MINERU_OCR_INSTALL_URL}"
    )


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert law text files (.txt/.docx/.pdf) to Markdown.")
    parser.add_argument("input", help="Input file path (.txt/.docx/.pdf)")
    parser.add_argument("--out", help="Output .md file path (overrides --out-dir)")
    parser.add_argument(
        "--out-dir",
        help='Output directory (default: create "markdown/" under input file directory)',
    )
    parser.add_argument(
        "--allow-fallback",
        action="store_true",
        help="Allow local fallback engines when mineru-ocr is unavailable or fails.",
    )
    parser.add_argument(
        "--docx-engine",
        choices=["auto", "mineru", "python-docx"],
        default="auto",
        help="DOCX extraction engine (default: auto = mineru-ocr skill; fallback needs --allow-fallback)",
    )
    parser.add_argument(
        "--pdf-engine",
        choices=["auto", "mineru", "pdfplumber"],
        default="auto",
        help="PDF extraction engine (default: auto = mineru-ocr skill; fallback needs --allow-fallback)",
    )
    parser.add_argument(
        "--skip-mineru-ocr-skill",
        action="store_true",
        help="Skip calling installed mineru-ocr skill and use local fallback engines.",
    )
    parser.add_argument(
        "--law-decision",
        choices=["auto", "law", "non-law"],
        default="auto",
        help="Stage2 document type decision from caller (auto=rule fallback).",
    )
    parser.add_argument(
        "--skip-stage3-check",
        action="store_true",
        help="Skip Stage3 checking/report generation.",
    )
    parser.add_argument(
        "--stage3-max-retries",
        type=int,
        default=2,
        help="Maximum Stage3-driven retry count (default: 2).",
    )
    parser.add_argument(
        "--stage3-strict",
        dest="stage3_strict",
        action="store_true",
        default=True,
        help="Fail the run when Stage3 does not pass after retries.",
    )
    parser.add_argument(
        "--no-stage3-strict",
        dest="stage3_strict",
        action="store_false",
        help="Do not fail the run when Stage3 fails.",
    )
    parser.add_argument(
        "--artifact-level",
        choices=ARTIFACT_LEVELS,
        default="minimal",
        help="Output artifact level: minimal|standard|debug (default: minimal).",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()

    parsed = urlparse(args.input)
    input_is_url = parsed.scheme in {"http", "https"} and bool(parsed.netloc)

    input_path: Path | None
    if input_is_url:
        input_path = None
    else:
        input_path = Path(args.input).expanduser().resolve()
        if not input_path.exists():
            raise SystemExit(f"Input not found: {input_path}")

    if args.out and args.out_dir:
        raise SystemExit("Use either --out or --out-dir, not both.")

    input_name = (
        Path(parsed.path).name if input_is_url and parsed.path else (input_path.name if input_path else "input")
    )
    input_stem = Path(input_name).stem or "output"

    if args.out:
        out_base_path = Path(args.out).expanduser().resolve()
    else:
        if args.out_dir:
            out_dir = Path(args.out_dir).expanduser().resolve()
        else:
            base_dir = input_path.parent if input_path else Path.cwd()
            out_dir = base_dir / "markdown"
        out_file_dir = out_dir / input_stem
        out_base_path = out_file_dir / f"{input_stem}.md"

    stage1_path, stage2_path = _derive_stage_paths(out_base_path)

    suffix = Path(parsed.path).suffix.lower() if input_is_url else input_path.suffix.lower()

    if suffix == ".txt":
        if input_is_url:
            raise SystemExit("For .txt, please provide a local file path (URL input not supported for txt).")
        assert input_path is not None
        _convert_txt(input_path, stage1_path)
        _prepare_stage2(stage1_path, stage2_path)
        pipeline = _run_stage2_stage3_pipeline(
            stage1_path=stage1_path,
            stage2_path=stage2_path,
            law_decision=args.law_decision,
            skip_stage3_check=args.skip_stage3_check,
            stage3_max_retries=args.stage3_max_retries,
        )
        _finalize_conversion(
            input_ref=str(input_path),
            input_stem=input_stem,
            stage1_path=stage1_path,
            stage2_path=stage2_path,
            pipeline=pipeline,
            law_decision=args.law_decision,
            engine="txt-copy",
            artifact_level=args.artifact_level,
            stage3_strict=args.stage3_strict,
            stage1_log_label="Saved stage1",
        )
        return

    if suffix == ".docx":
        if input_is_url:
            raise SystemExit("DOCX currently supports local files only when using mineru-ocr skill.")
        assert input_path is not None

        skill_error = ""
        if not args.skip_mineru_ocr_skill and args.docx_engine in {"auto", "mineru"}:
            ok, msg = _convert_with_mineru_ocr_skill(input_path, stage1_path)
            if ok:
                _prepare_stage2(stage1_path, stage2_path)
                pipeline = _run_stage2_stage3_pipeline(
                    stage1_path=stage1_path,
                    stage2_path=stage2_path,
                    law_decision=args.law_decision,
                    skip_stage3_check=args.skip_stage3_check,
                    stage3_max_retries=args.stage3_max_retries,
                )
                _finalize_conversion(
                    input_ref=str(input_path),
                    input_stem=input_stem,
                    stage1_path=stage1_path,
                    stage2_path=stage2_path,
                    pipeline=pipeline,
                    law_decision=args.law_decision,
                    engine="mineru-ocr-skill",
                    artifact_level=args.artifact_level,
                    stage3_strict=args.stage3_strict,
                    stage1_log_label="Saved stage1",
                    print_engine="Engine: mineru-ocr-skill",
                )
                return
            skill_error = msg

        if args.docx_engine == "python-docx" or args.allow_fallback:
            _convert_docx(input_path, stage1_path)
            _prepare_stage2(stage1_path, stage2_path)
            pipeline = _run_stage2_stage3_pipeline(
                stage1_path=stage1_path,
                stage2_path=stage2_path,
                law_decision=args.law_decision,
                skip_stage3_check=args.skip_stage3_check,
                stage3_max_retries=args.stage3_max_retries,
            )
            _finalize_conversion(
                input_ref=str(input_path),
                input_stem=input_stem,
                stage1_path=stage1_path,
                stage2_path=stage2_path,
                pipeline=pipeline,
                law_decision=args.law_decision,
                engine="python-docx-fallback",
                artifact_level=args.artifact_level,
                stage3_strict=args.stage3_strict,
                stage1_log_label="Saved stage1 (fallback)",
            )
            return

        if args.skip_mineru_ocr_skill:
            raise SystemExit(
                "mineru-ocr skill is skipped by --skip-mineru-ocr-skill. "
                "Use --allow-fallback to continue with python-docx."
            )

        if "mineru-ocr skill not found" in skill_error:
            raise SystemExit(
                "mineru-ocr skill failed for .docx.\n"
                f"{_mineru_ocr_install_hint()}\n"
                "或使用 --allow-fallback 走 python-docx 回退。"
            )

        raise SystemExit(
            "mineru-ocr skill failed for .docx. "
            "Check mineru-ocr config/token, "
            "or re-run with --allow-fallback.\n"
            f"- mineru-ocr error: {skill_error}"
        )

    if suffix == ".pdf":
        if input_is_url:
            raise SystemExit("PDF currently supports local files only when using mineru-ocr skill.")
        assert input_path is not None

        skill_error = ""
        if not args.skip_mineru_ocr_skill and args.pdf_engine in {"auto", "mineru"}:
            ok, msg = _convert_with_mineru_ocr_skill(input_path, stage1_path)
            if ok:
                _prepare_stage2(stage1_path, stage2_path)
                pipeline = _run_stage2_stage3_pipeline(
                    stage1_path=stage1_path,
                    stage2_path=stage2_path,
                    law_decision=args.law_decision,
                    skip_stage3_check=args.skip_stage3_check,
                    stage3_max_retries=args.stage3_max_retries,
                )
                _finalize_conversion(
                    input_ref=str(input_path),
                    input_stem=input_stem,
                    stage1_path=stage1_path,
                    stage2_path=stage2_path,
                    pipeline=pipeline,
                    law_decision=args.law_decision,
                    engine="mineru-ocr-skill",
                    artifact_level=args.artifact_level,
                    stage3_strict=args.stage3_strict,
                    stage1_log_label="Saved stage1",
                    print_engine="Engine: mineru-ocr-skill",
                )
                return
            skill_error = msg

        if args.pdf_engine == "pdfplumber" or args.allow_fallback:
            _convert_pdf_pdfplumber(input_path, stage1_path)
            _prepare_stage2(stage1_path, stage2_path)
            pipeline = _run_stage2_stage3_pipeline(
                stage1_path=stage1_path,
                stage2_path=stage2_path,
                law_decision=args.law_decision,
                skip_stage3_check=args.skip_stage3_check,
                stage3_max_retries=args.stage3_max_retries,
            )
            _finalize_conversion(
                input_ref=str(input_path),
                input_stem=input_stem,
                stage1_path=stage1_path,
                stage2_path=stage2_path,
                pipeline=pipeline,
                law_decision=args.law_decision,
                engine="pdfplumber-fallback",
                artifact_level=args.artifact_level,
                stage3_strict=args.stage3_strict,
                stage1_log_label="Saved stage1 (fallback)",
            )
            return

        if args.skip_mineru_ocr_skill:
            raise SystemExit(
                "mineru-ocr skill is skipped by --skip-mineru-ocr-skill. "
                "Use --allow-fallback to continue with pdfplumber."
            )

        if "mineru-ocr skill not found" in skill_error:
            raise SystemExit(
                "mineru-ocr skill failed for .pdf.\n"
                f"{_mineru_ocr_install_hint()}\n"
                "或使用 --allow-fallback 走 pdfplumber 回退。"
            )

        raise SystemExit(
            "mineru-ocr skill failed for .pdf. "
            "Check mineru-ocr config/token, "
            "or re-run with --allow-fallback.\n"
            f"- mineru-ocr error: {skill_error}"
        )

    raise SystemExit(f"Unsupported file type: {input_name}")


if __name__ == "__main__":
    main()
